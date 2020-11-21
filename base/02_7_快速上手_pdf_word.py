#!/usr/bin/python
# -*- coding:utf-8 -*-
# 第二部分 自动化任务

print('********* 第 13 章 处理 PDF 和 Word 文档 ***********')

import PyPDF2

# 从 pdf 提取文本
# pdfFileObj = open('data\\meetingminutes.pdf', 'rb')
# pdfReader = PyPDF2.PdfFileReader(pdfFileObj)
# print(pdfReader.numPages)  # 19
# pageObj = pdfReader.getPage(0)
# print(pageObj.extractText())

# 解密 PDF encrypted.pdf
# pdfFileObj = open('data\\encrypted.pdf', 'rb')
# pdfReader = PyPDF2.PdfFileReader(pdfFileObj)
# print(pdfReader.isEncrypted)  # True
# # pdfReader.getPage(0)  # 报错
# pdfReader.decrypt('rosebud')  # 解密
# print(pdfReader.getPage(0))

# 拷贝 PDF
# pdf1File = open('data\\meetingminutes.pdf', 'rb')
# pdf2File = open('data\\meetingminutes2.pdf', 'rb')
# pdf1Reader = PyPDF2.PdfFileReader(pdf1File)
# pdf2Reader = PyPDF2.PdfFileReader(pdf2File)
# pdfWriter = PyPDF2.PdfFileWriter()
# for pageNum in range(pdf1Reader.numPages):
#     pageObj = pdf1Reader.getPage(pageNum)
#     pdfWriter.addPage(pageObj)
# for pageNum in range(pdf2Reader.numPages):
#     pageObj = pdf2Reader.getPage(pageNum)
#     pdfWriter.addPage(pageObj)
# pdfOutputFile = open('data\\combinedminutes.pdf', 'wb')
# pdfWriter.write(pdfOutputFile)
# pdfOutputFile.close()
# pdf1File.close()
# pdf2File.close()

# 旋转页面
# minutesFile = open('data\\meetingminutes.pdf', 'rb')
# pdfReader = PyPDF2.PdfFileReader(minutesFile)
# page = pdfReader.getPage(0)
# page.rotateClockwise(90)
# pdfWriter = PyPDF2.PdfFileWriter()
# pdfWriter.addPage(page)
# resultPdfFile = open('data\\rotatedPage.pdf', 'wb')
# pdfWriter.write(resultPdfFile)
# resultPdfFile.close()
# minutesFile.close()

# 叠加页面，加水印
# minutesFile = open('data\\meetingminutes.pdf', 'rb')
# pdfReader = PyPDF2.PdfFileReader(minutesFile)
# minutesFirstPage = pdfReader.getPage(0)
# pdfWatermarkReader = PyPDF2.PdfFileReader(open('data\\watermark.pdf', 'rb'))
# minutesFirstPage.mergePage(pdfWatermarkReader.getPage(0))
# pdfWriter = PyPDF2.PdfFileWriter()
# pdfWriter.addPage(minutesFirstPage)
# for pageNum in range(1, pdfReader.numPages):
#     pageObj = pdfReader.getPage(pageNum)
#     pdfWriter.addPage(pageObj)
# resultPdfFile = open('data\\watermarkedCover.pdf', 'wb')
# pdfWriter.write(resultPdfFile)
# minutesFile.close()
# resultPdfFile.close()

# 加密 PDF
# pdfFile = open('data\\meetingminutes.pdf', 'rb')
# pdfReader = PyPDF2.PdfFileReader(pdfFile)
# pdfWriter = PyPDF2.PdfFileWriter()
# for pageNum in range(pdfReader.numPages):
#     pdfWriter.addPage(pdfReader.getPage(pageNum))
# pdfWriter.encrypt('swordfish')
# resultPdf = open('data\\encryptedminutes.pdf', 'wb')
# pdfWriter.write(resultPdf)
# resultPdf.close()
# pdfFile.close()

import docx

# 读取 Word 文档
# doc = docx.Document('data\\demo.docx')
# print(len(doc.paragraphs))  # 7
# print(doc.paragraphs[0].text)  # Document Title
# print(doc.paragraphs[1].text)  # A plain paragraph with some bold and some italic
# print(len(doc.paragraphs[1].runs))  # 5
# print(doc.paragraphs[1].runs[0].text)  # A plain paragraph with
# print(doc.paragraphs[1].runs[1].text)  # some
# print(doc.paragraphs[1].runs[2].text)  # bold
# print(doc.paragraphs[1].runs[3].text)  # and some
# print(doc.paragraphs[1].runs[4].text)  # italic

# 读取完整的文本
# def getText(filename):
#     doc = docx.Document(filename)
#     fullText = []
#     for para in doc.paragraphs:
#         fullText.append(para.text)
#     return '\n'.join(fullText)
# print(getText('data\\demo.docx'))

# 写入 Word 文档
# doc = docx.Document()
# print(doc.add_paragraph('Hello world!'))
# paraObj1 = doc.add_paragraph('This is a second paragraph.')
# paraObj2 = doc.add_paragraph('This is a yet another paragraph.')
# paraObj1.add_run(' This text is being added to the second paragraph.')
# # doc.save('data\\helloworld.docx')
# doc.save('data\\multipleParagraphs.docx')




